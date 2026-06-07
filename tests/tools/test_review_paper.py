import sys
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import docx as docx_lib
import pytest
import pytest_asyncio

sys.path.insert(0, str(Path(__file__).parent.parent.parent))


def test_load_api_key_reads_and_strips(tmp_path):
    key_file = tmp_path / "api_key.txt"
    key_file.write_text("  sk-ant-test123\n  ")

    from tools.review_paper import load_api_key
    with patch("tools.review_paper.API_KEY_PATH", key_file):
        key = load_api_key()

    assert key == "sk-ant-test123"


def test_extract_docx_text_includes_headings_and_body(tmp_path):
    doc = docx_lib.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the intro text.")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("We used PatchTST.")
    doc_path = tmp_path / "test.docx"
    doc.save(str(doc_path))

    from tools.review_paper import extract_docx_text
    text = extract_docx_text(doc_path)

    assert "Introduction" in text
    assert "This is the intro text." in text
    assert "Methods" in text
    assert "We used PatchTST." in text


def test_extract_docx_text_nonexistent_raises(tmp_path):
    from tools.review_paper import extract_docx_text
    with pytest.raises(Exception):
        extract_docx_text(tmp_path / "missing.docx")


def test_extract_pdf_text_nonexistent_raises(tmp_path):
    from tools.review_paper import extract_pdf_text
    with pytest.raises(FileNotFoundError):
        extract_pdf_text(tmp_path / "missing.pdf")


@pytest.mark.asyncio
async def test_run_agent_writes_report_to_file(tmp_path):
    from tools.review_paper import run_agent

    mock_client = MagicMock()
    mock_message = MagicMock()
    mock_content = MagicMock(text="# ScientificReviewer Review\n\n## Summary\nSolid.")
    mock_content.type = "text"
    mock_message.content = [mock_content]
    mock_client.messages.create = AsyncMock(return_value=mock_message)

    result = await run_agent(mock_client, "scientific", "paper text here", "", tmp_path)

    assert "ScientificReviewer" in result
    report_path = tmp_path / "scientific.md"
    assert report_path.exists()
    assert report_path.read_text() == result


@pytest.mark.asyncio
async def test_run_agent_proofreader_sends_springer_text(tmp_path):
    from tools.review_paper import run_agent

    mock_client = MagicMock()
    mock_message = MagicMock()
    mock_content = MagicMock(text="# Proofreader Review\n\n## Summary\nOK.")
    mock_content.type = "text"
    mock_message.content = [mock_content]
    mock_client.messages.create = AsyncMock(return_value=mock_message)

    await run_agent(mock_client, "proofread", "paper text", "springer instructions", tmp_path)

    call_kwargs = mock_client.messages.create.call_args.kwargs
    user_content = call_kwargs["messages"][0]["content"]
    assert "springer instructions" in user_content


@pytest.mark.asyncio
async def test_run_agent_scientific_does_not_send_springer_text(tmp_path):
    from tools.review_paper import run_agent

    mock_client = MagicMock()
    mock_message = MagicMock()
    mock_content = MagicMock(text="# ScientificReviewer Review\n\n## Summary\nOK.")
    mock_content.type = "text"
    mock_message.content = [mock_content]
    mock_client.messages.create = AsyncMock(return_value=mock_message)

    await run_agent(mock_client, "scientific", "paper text", "springer instructions", tmp_path)

    call_kwargs = mock_client.messages.create.call_args.kwargs
    user_content = call_kwargs["messages"][0]["content"]
    assert "springer instructions" not in user_content


@pytest.mark.asyncio
async def test_run_orchestrator_writes_synthesis(tmp_path):
    from tools.review_paper import run_orchestrator

    mock_client = MagicMock()
    mock_message = MagicMock()
    mock_content = MagicMock()
    mock_content.type = "text"
    mock_content.text = "# Paper Review Synthesis\n\n## Overall Assessment\nReady to submit."
    mock_message.content = [mock_content]
    mock_client.messages.create = AsyncMock(return_value=mock_message)

    reports = {
        "ScientificReviewer": "sci report text",
        "RhetoricReviewer": "rhetoric report text",
    }
    await run_orchestrator(mock_client, reports, tmp_path)

    synthesis_path = tmp_path / "synthesis.md"
    assert synthesis_path.exists()
    assert "Synthesis" in synthesis_path.read_text()


@pytest.mark.asyncio
async def test_run_orchestrator_includes_all_reports_in_prompt(tmp_path):
    from tools.review_paper import run_orchestrator

    mock_client = MagicMock()
    mock_message = MagicMock()
    mock_content = MagicMock()
    mock_content.type = "text"
    mock_content.text = "# Paper Review Synthesis\n\nDone."
    mock_message.content = [mock_content]
    mock_client.messages.create = AsyncMock(return_value=mock_message)

    reports = {
        "ScientificReviewer": "unique-sci-text",
        "Proofreader": "unique-proof-text",
    }
    await run_orchestrator(mock_client, reports, tmp_path)

    call_kwargs = mock_client.messages.create.call_args.kwargs
    user_content = call_kwargs["messages"][0]["content"]
    assert "unique-sci-text" in user_content
    assert "unique-proof-text" in user_content


@pytest.mark.asyncio
async def test_run_review_creates_all_output_files(tmp_path):
    from tools.review_paper import run_review

    # Create a minimal docx
    doc = docx_lib.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Some paper content here.")
    paper_path = tmp_path / "paper.docx"
    doc.save(str(paper_path))

    # Mock api key
    key_path = tmp_path / "api_key.txt"
    key_path.write_text("sk-fake-key")

    def make_mock_message(text):
        msg = MagicMock()
        mock_content = MagicMock()
        mock_content.type = "text"
        mock_content.text = text
        msg.content = [mock_content]
        return msg

    call_count = 0

    async def fake_create(**kwargs):
        nonlocal call_count
        call_count += 1
        # First 4 calls = specialists, 5th = orchestrator
        if call_count <= 4:
            system = kwargs.get("system", "")
            name = "ScientificReviewer" if "NeurIPS" in system else \
                   "RhetoricReviewer" if "rhetoric" in system.lower() else \
                   "Proofreader" if "copy editor" in system.lower() else \
                   "AIDetectionReviewer"
            return make_mock_message(f"# {name} Review\n\n## Summary\nTest.")
        return make_mock_message("# Paper Review Synthesis\n\n## Overall Assessment\nOK.")

    mock_client = MagicMock()
    mock_client.messages.create = AsyncMock(side_effect=fake_create)

    with patch("tools.review_paper.API_KEY_PATH", key_path), \
         patch("tools.review_paper.anthropic.AsyncAnthropic", return_value=mock_client), \
         patch("tools.review_paper.SPRINGER_PDF_PATH", tmp_path / "missing.pdf"):

        out_dir = await run_review(
            paper_path=paper_path,
            out_base=tmp_path / "review",
        )

    assert (out_dir / "synthesis.md").exists()
    existing = list(out_dir.glob("*.md"))
    assert len(existing) == 5  # 4 specialists + synthesis
