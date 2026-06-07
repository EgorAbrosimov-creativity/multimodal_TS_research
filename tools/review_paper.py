import asyncio
import argparse
import os
import sys
from datetime import date
from pathlib import Path

import anthropic
import docx
import pypdf

ROOT = Path(__file__).parent.parent
API_KEY_PATH = ROOT / ".claude" / "api_key.txt"
SPRINGER_PDF_PATH = ROOT / "docs" / "Instructions+for+proceedings+authors+(pdf).pdf"
DEFAULT_PAPER = ROOT / "paper_draft.docx"
DEFAULT_OUT = ROOT / "docs" / "review"
MODEL = "claude-opus-4-7"
TODAY = str(date.today())


def load_api_key() -> str:
    if env_key := os.environ.get("ANTHROPIC_API_KEY"):
        return env_key
    return API_KEY_PATH.read_text().strip()


def extract_docx_text(path: Path) -> str:
    doc = docx.Document(str(path))
    lines = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if para.style.name.startswith("Heading"):
            raw_level = para.style.name.split()[-1]
            level = int(raw_level) if raw_level.isdigit() else 1
            lines.append(f"\n{'#' * level} {para.text}\n")
        else:
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)
    return "\n".join(lines)


def extract_pdf_text(path: Path) -> str:
    reader = pypdf.PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


_REPORT_FORMAT = """
Output a structured markdown report with this exact format:

# {name} Review — paper_draft.docx
*Reviewed: {today}*

## Summary
(2–3 sentences overall assessment)

## Critical Issues
(Each as: - **[Section N]** Description. Quote problematic text if relevant.)

## Major Issues
(Same format)

## Minor Issues
(Same format)

## Positive Notes
(Briefly note genuine strengths)
"""


def _report_fmt(name: str) -> str:
    return _REPORT_FORMAT.replace("{name}", name).replace("{today}", TODAY)


AGENTS: dict[str, dict] = {
    "scientific": {
        "name": (_n := "ScientificReviewer"),
        "filename": "scientific.md",
        "needs_springer": False,
        "system": (
            "You are a senior ML conference reviewer with expertise in time series forecasting "
            "and multimodal learning (NeurIPS/ICML style).\n\n"
            "Review the provided academic paper. Focus exclusively on scientific rigor:\n"
            "- Verify each contribution claim in §1 Introduction is supported by experimental "
            "results in §6–7\n"
            "- Check all numbers in tables for internal consistency (degradation ratios vs raw "
            "MSE values, means vs reported conclusions)\n"
            "- Flag unsupported or overstated claims (e.g. 'outperforms' without statistical "
            "backing, causal claims from correlational evidence)\n"
            "- Identify missing, misrepresented, or insufficient citations in §3 Related Work\n"
            "- Assess whether the failure mode taxonomy (α-collapse, gate suppression, FiLM "
            "fragility) is adequately evidenced by the described diagnostics\n"
            "- Check experimental protocol completeness: seeds reported, train/val/test splits "
            "defined, metrics specified, datasets described\n\n"
        ) + _report_fmt(_n),
    },
    "rhetoric": {
        "name": (_n := "RhetoricReviewer"),
        "filename": "rhetoric.md",
        "needs_springer": False,
        "system": (
            "You are an academic writing coach with experience in machine learning venues.\n\n"
            "Review the provided paper for argument structure and rhetoric only. Do not comment "
            "on grammar, spelling, or formatting. Focus exclusively on:\n"
            "- Whether the abstract accurately reflects the body (contributions, key findings, "
            "datasets used, main conclusions)\n"
            "- Intro↔conclusion alignment: does the conclusion answer what the introduction "
            "promised?\n"
            "- Clarity and parallelism of the four stated contributions\n"
            "- Section-by-section argument flow: does each section set up the next?\n"
            "- Missing or weak transitions between sections\n"
            "- Sections disproportionately long or short relative to their role\n"
            "- Whether §8 Practitioner Guide follows naturally from §6–7 results\n"
            "- Whether the revised framing ('which fusion mechanisms can exploit text') is "
            "consistently maintained throughout\n\n"
        ) + _report_fmt(_n),
    },
    "proofread": {
        "name": (_n := "Proofreader"),
        "filename": "proofread.md",
        "needs_springer": True,
        "system": (
            "You are an expert copy editor with academic publishing experience, specialising "
            "in Springer LNCS/CCIS proceedings.\n\n"
            "Review the paper for language errors AND Springer compliance. You will receive "
            "the paper text and the official Springer author instructions.\n\n"
            "LANGUAGE (check every section):\n"
            "- Grammar, spelling, and punctuation errors (quote the error and give the correction)\n"
            "- British/American English consistency — identify which is used, flag deviations\n"
            "- Number formatting: '%' vs 'percent' (pick one), spacing around '=', decimal "
            "separators\n"
            "- Hyphenation consistency (e.g. 'text-augmented' must be hyphenated consistently)\n"
            "- Sentence-level clarity issues: ambiguous pronouns, dangling modifiers\n\n"
            "SPRINGER LNCS/CCIS COMPLIANCE:\n"
            "- Heading capitalisation: nouns, verbs, adjectives capitalised; articles, "
            "prepositions, conjunctions lowercase — check every heading\n"
            "- Only H1 and H2 numbered; H3 and below unnumbered\n"
            "- Citations in brackets [n], never superscript; multiple citations as [4-6, 9] "
            "in numerical order\n"
            "- Figure captions below figures; table captions above tables\n"
            "- Acknowledgments section present (third-level heading)\n"
            "- Disclosure of Interests section present\n"
            "- Corresponding author marked with a symbol in the header; email mandatory\n"
            "- ORCID identifiers in header\n"
            "- Full paper target: 12–15 pages\n\n"
            "For each issue quote the original text and give the corrected version.\n\n"
        ) + _report_fmt(_n),
    },
    "ai_detection": {
        "name": (_n := "AIDetectionReviewer"),
        "filename": "ai_detection.md",
        "needs_springer": False,
        "system": (
            "You are an AI writing naturalness auditor. Identify text that reads as "
            "AI-generated and propose specific natural rewrites.\n\n"
            "Focus on these patterns in priority order:\n"
            "1. Vague intensifiers: 'robust', 'comprehensive', 'significant', 'novel', "
            "'noteworthy', 'state-of-the-art' used without specific referents\n"
            "2. Over-hedging: 'it is worth noting that', 'it can be observed that', "
            "'it is important to note', 'in this regard'\n"
            "3. Hollow section openers: 'In this section, we...', 'We now turn to...', "
            "'Having established X, we now...'\n"
            "4. Repetitive sentence structures: consecutive sentences beginning with 'We' or "
            "following the same Subject-Verb-Object pattern\n"
            "5. LLM-signature phrases: 'delve into', 'shed light on', 'testament to', "
            "'in the realm of', 'a myriad of', 'the intersection of'\n"
            "6. Empty transition summaries: 'As discussed above...', 'As shown in the "
            "previous section...'\n\n"
            "Check abstract, introduction, and conclusion with highest priority.\n\n"
            "For EACH flagged passage:\n"
            "1. Give the section reference (e.g. §1, Abstract)\n"
            "2. Quote the exact original text\n"
            "3. Explain in one sentence what makes it sound AI-generated\n"
            "4. Propose a specific, natural rewrite\n\n"
            "Do not flag standard technical or genuinely appropriate academic phrasing.\n\n"
        ) + _report_fmt(_n),
    },
}

ORCHESTRATOR_SYSTEM = (
    "You are a review coordinator synthesising four specialist reviews of an academic paper.\n\n"
    "Your job:\n"
    "1. Read all four reports (ScientificReviewer, RhetoricReviewer, Proofreader, "
    "AIDetectionReviewer)\n"
    "2. Identify duplicate or overlapping findings — merge into a single item, noting all "
    "source agents\n"
    "3. Assign severity:\n"
    "   - Critical: blocks submission (factual errors, unsupported central claims, major "
    "Springer violations)\n"
    "   - Major: strongly recommended before submission\n"
    "   - Minor: polish (typos, minor phrasing, small inconsistencies)\n"
    "4. Produce a single ranked action list\n\n"
    "Output this exact format:\n\n"
    "# Paper Review Synthesis\n"
    f"*Generated: {TODAY}*\n\n"
    "## Overall Assessment\n"
    "(One paragraph: paper's readiness for submission, strongest areas, most urgent gaps)\n\n"
    "## Critical — fix before submission\n"
    "- **[SourceAgent(s)]** Description of issue and specific fix.\n\n"
    "## Major — strongly recommended\n"
    "- **[SourceAgent(s)]** ...\n\n"
    "## Minor — polish\n"
    "- **[SourceAgent(s)]** ...\n\n"
    "## Duplicate findings resolved\n"
    "(Brief note on findings that appeared in multiple reports and how they were merged)\n"
)


async def run_agent(
    client: anthropic.AsyncAnthropic,
    agent_key: str,
    paper_text: str,
    springer_text: str,
    out_dir: Path,
) -> str:
    if agent_key not in AGENTS:
        raise ValueError(
            f"Unknown agent_key {agent_key!r}. Valid keys: {list(AGENTS)}"
        )
    agent = AGENTS[agent_key]
    user_content = f"<paper>\n{paper_text}\n</paper>"
    if agent["needs_springer"]:
        user_content += f"\n\n<springer_instructions>\n{springer_text}\n</springer_instructions>"
    user_content += f"\n\nToday's date: {TODAY}"

    try:
        message = await client.messages.create(
            model=MODEL,
            max_tokens=4096,
            system=agent["system"],
            messages=[{"role": "user", "content": user_content}],
        )
    except anthropic.APIError as exc:
        raise RuntimeError(f"[{agent['name']}] API error: {exc}") from exc
    if not message.content or message.content[0].type != "text":
        raise RuntimeError(
            f"[{agent['name']}] Unexpected API response: "
            f"content={message.content!r}"
        )
    report = message.content[0].text
    out_path = out_dir / agent["filename"]
    out_path.write_text(report, encoding="utf-8")
    print(f"[✓] {agent['name']} done → {out_path}")
    return report


async def run_orchestrator(
    client: anthropic.AsyncAnthropic,
    reports: dict[str, str],
    out_dir: Path,
) -> None:
    if not reports:
        raise ValueError("reports must not be empty")
    reports_text = "\n\n---\n\n".join(
        f"# {name} Report\n\n{text}" for name, text in reports.items()
    )
    try:
        message = await client.messages.create(
            model=MODEL,
            max_tokens=4096,
            system=ORCHESTRATOR_SYSTEM,
            messages=[{"role": "user", "content": f"<reports>\n{reports_text}\n</reports>"}],
        )
    except anthropic.APIError as exc:
        raise RuntimeError(f"[Orchestrator] API error: {exc}") from exc
    if not message.content or message.content[0].type != "text":
        raise RuntimeError(
            f"[Orchestrator] Unexpected API response: content={message.content!r}"
        )
    synthesis = message.content[0].text
    out_path = out_dir / "synthesis.md"
    out_path.write_text(synthesis, encoding="utf-8")
    print(f"[✓] Orchestrator done → {out_path}")


async def run_review(
    paper_path: Path = DEFAULT_PAPER,
    out_base: Path = DEFAULT_OUT,
    springer_path: Path = SPRINGER_PDF_PATH,
) -> Path:
    api_key = load_api_key()
    client = anthropic.AsyncAnthropic(api_key=api_key)

    print("Extracting paper text...")
    paper_text = extract_docx_text(paper_path)

    print("Extracting Springer instructions...")
    try:
        springer_text = extract_pdf_text(springer_path)
    except Exception as e:
        print(f"[!] Could not read Springer PDF: {e}. Proofreader runs without compliance context.")
        springer_text = ""

    out_dir = out_base / str(date.today())
    out_dir.mkdir(parents=True, exist_ok=True)
    print(f"Output directory: {out_dir}\n")

    print("Running specialist reviewers in parallel...")
    tasks = [
        run_agent(client, key, paper_text, springer_text, out_dir)
        for key in AGENTS
    ]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    reports: dict[str, str] = {}
    for key, result in zip(AGENTS.keys(), results):
        agent_name = AGENTS[key]["name"]
        if isinstance(result, Exception):
            print(f"[✗] {agent_name} failed: {result}")
            error_path = out_dir / AGENTS[key]["filename"]
            error_path.write_text(f"# ERROR\n\nAgent failed: {result}\n", encoding="utf-8")
        else:
            reports[agent_name] = result  # type: ignore[assignment]

    if not reports:
        print(f"[✗] All specialist agents failed. Check error stubs in {out_dir}")
        return out_dir

    print(f"Specialist reports written to: {out_dir}")
    print("\nRunning Orchestrator...")
    await run_orchestrator(client, reports, out_dir)
    return out_dir


def main() -> None:
    parser = argparse.ArgumentParser(description="Run BMAD-style paper review workflow")
    parser.add_argument("--paper", type=Path, default=DEFAULT_PAPER, help="Path to paper_draft.docx")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT, help="Base output directory")
    parser.add_argument("--springer-pdf", type=Path, default=SPRINGER_PDF_PATH,
                        help="Path to Springer author instructions PDF")
    args = parser.parse_args()
    try:
        asyncio.run(run_review(args.paper, args.out, args.springer_pdf))
    except Exception as exc:
        print(f"[✗] Review failed: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
